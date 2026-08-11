from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import tempfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


CITATION_RE = re.compile(r"\[(?:\d+)(?:(?:-|,|，)\d+)*\]")
REFERENCE_HEADING_RE = re.compile(r"^\s*参考文献\s*$")


def reference_start(document):
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "参考文献标题" or REFERENCE_HEADING_RE.match(paragraph.text):
            return index
    return len(document.paragraphs)


def is_pure_text_run(run_element):
    allowed = {qn("w:rPr"), qn("w:t")}
    return all(child.tag in allowed for child in run_element)


def set_superscript(run_element):
    rpr = run_element.get_or_add_rPr()
    vert_align = rpr.find(qn("w:vertAlign"))
    if vert_align is None:
        vert_align = OxmlElement("w:vertAlign")
        rpr.append(vert_align)
    vert_align.set(qn("w:val"), "superscript")


def make_text_run(source, text, superscript):
    clone = deepcopy(source)
    for child in list(clone):
        if child.tag != qn("w:rPr"):
            clone.remove(child)
    text_node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    clone.append(text_node)
    if superscript:
        set_superscript(clone)
    return clone


def apply_to_paragraph(paragraph, paragraph_index):
    text = paragraph.text
    intervals = [(m.start(), m.end(), m.group(0)) for m in CITATION_RE.finditer(text)]
    if not intervals:
        return []

    run_positions = []
    position = 0
    for run in paragraph.runs:
        start, end = position, position + len(run.text)
        run_positions.append((run, start, end))
        position = end

    for run, run_start, run_end in run_positions:
        overlaps = []
        for match_start, match_end, token in intervals:
            overlap_start = max(match_start, run_start)
            overlap_end = min(match_end, run_end)
            if overlap_start < overlap_end:
                overlaps.append((overlap_start - run_start, overlap_end - run_start, token))
        if not overlaps:
            continue
        if not is_pure_text_run(run._r):
            tokens = "、".join(dict.fromkeys(item[2] for item in overlaps))
            raise RuntimeError(
                f"段落 {paragraph_index} 的引用 {tokens} 与字段、图片或特殊对象共用 run，已停止以避免破坏文档。"
            )
        boundaries = {0, len(run.text)}
        for start, end, _ in overlaps:
            boundaries.update((start, end))
        ordered = sorted(boundaries)
        pieces = []
        for start, end in zip(ordered, ordered[1:]):
            segment = run.text[start:end]
            if not segment:
                continue
            superscript = any(start >= cite_start and end <= cite_end for cite_start, cite_end, _ in overlaps)
            pieces.append(make_text_run(run._r, segment, superscript))
        parent = run._r.getparent()
        insertion_index = parent.index(run._r)
        parent.remove(run._r)
        for offset, piece in enumerate(pieces):
            parent.insert(insertion_index + offset, piece)

    return [{"paragraph_index": paragraph_index, "citation": token} for _, _, token in intervals]


def main():
    parser = argparse.ArgumentParser(description="将论文正文中的方括号数字引用设为上角标。")
    parser.add_argument("docx", type=Path)
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--in-place", action="store_true", help="安全替换原 DOCX，并在系统临时目录备份")
    group.add_argument("--output", type=Path, help="写入另一个 DOCX")
    parser.add_argument("--report", type=Path, help="写入 JSON 修改报告")
    args = parser.parse_args()

    source = args.docx.resolve()
    if not source.is_file() or source.suffix.lower() != ".docx":
        parser.error("输入必须是存在的 .docx 文件")

    document = Document(source)
    before = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(document.inline_shapes),
        "sections": len(document.sections),
    }
    ref_start = reference_start(document)
    changes = []
    for index, paragraph in enumerate(document.paragraphs[:ref_start]):
        changes.extend(apply_to_paragraph(paragraph, index))

    backup = None
    if args.in_place:
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        backup_dir = Path(tempfile.gettempdir()) / "thesis-format-normalize-backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        backup = backup_dir / f"{source.stem}-引用上角标前-{stamp}{source.suffix}"
        shutil.copy2(source, backup)
        fd, temp_name = tempfile.mkstemp(prefix=f".{source.stem}-", suffix=".docx", dir=source.parent)
        os.close(fd)
        temp_output = Path(temp_name)
        try:
            document.save(temp_output)
            os.replace(temp_output, source)
        finally:
            if temp_output.exists():
                temp_output.unlink()
        destination = source
    else:
        destination = args.output.resolve()
        destination.parent.mkdir(parents=True, exist_ok=True)
        document.save(destination)

    verified = Document(destination)
    after = {
        "paragraphs": len(verified.paragraphs),
        "tables": len(verified.tables),
        "images": len(verified.inline_shapes),
        "sections": len(verified.sections),
    }
    if before != after:
        raise RuntimeError(f"结构计数发生变化：修改前 {before}，修改后 {after}")

    result = {
        "document": str(destination),
        "backup": str(backup) if backup else None,
        "reference_start_paragraph": ref_start,
        "changed_citations": len(changes),
        "changes": changes,
        "structure_before": before,
        "structure_after": after,
    }
    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
