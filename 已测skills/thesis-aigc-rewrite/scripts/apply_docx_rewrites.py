from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import tempfile
from pathlib import Path

from docx import Document


NUMBER_RE = re.compile(r"(?<![A-Za-z_])\d+(?:\.\d+)*(?:/\d+(?:\.\d+)*)?")
LATIN_TOKEN_RE = re.compile(
    r"(?<![A-Za-z0-9_])(?:[A-Za-z][A-Za-z0-9_]*(?:[./:@=-][A-Za-z0-9_]+)+|"
    r"[A-Za-z]+\d+[A-Za-z0-9]*|[A-Z]{2,}[A-Za-z0-9_-]*)(?![A-Za-z0-9_])"
)
CITATION_RE = re.compile(r"\[\d+(?:[-,]\d+)*\]")


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text)


def protected(text: str) -> dict[str, list[str]]:
    return {
        "numbers": NUMBER_RE.findall(text),
        "technical_tokens": LATIN_TOKEN_RE.findall(text),
        "citations": CITATION_RE.findall(text),
    }


def structure(document: Document) -> dict[str, int]:
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(document.inline_shapes),
        "headings": sum(1 for p in document.paragraphs if p.style.name.startswith("Heading")),
        "sections": len(document.sections),
    }


def replace_text(paragraph, new_text: str) -> None:
    element = paragraph._p
    paragraph_properties = element.pPr
    for child in list(element):
        if child is not paragraph_properties:
            element.remove(child)
    paragraph.add_run(new_text)


def main() -> None:
    parser = argparse.ArgumentParser(description="Safely apply reviewed paragraph rewrites to one DOCX")
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--rewrites", type=Path, required=True)
    parser.add_argument("--audit", type=Path, required=True)
    args = parser.parse_args()

    items = json.loads(args.rewrites.read_text(encoding="utf-8-sig"))
    if not isinstance(items, list) or not items:
        raise ValueError("Rewrites JSON must be a non-empty list")

    document = Document(args.docx)
    before = structure(document)
    seen = set()
    audit_items = []

    for item in items:
        index = int(item["paragraph_index"])
        if index in seen:
            raise ValueError(f"Duplicate paragraph_index: {index}")
        seen.add(index)
        paragraph = document.paragraphs[index]
        current = paragraph.text.strip()
        expected = item["original_text"].strip()
        replacement = item["new_text"].strip()
        if normalize(current) != normalize(expected):
            raise RuntimeError(f"Source paragraph changed or index is wrong: {index}")
        old_protected = protected(current)
        new_protected = protected(replacement)
        missing = {
            category: sorted(set(values) - set(new_protected[category]))
            for category, values in old_protected.items()
            if set(values) - set(new_protected[category])
        }
        if missing:
            raise RuntimeError(f"Protected content missing at paragraph {index}: {missing}")
        audit_items.append({
            "paragraph_index": index,
            "style": paragraph.style.name,
            "original_text": current,
            "new_text": replacement,
            "protected": old_protected,
        })
        replace_text(paragraph, replacement)

    backup_dir = Path(tempfile.mkdtemp(prefix="codex_thesis_aigc_rewrite_"))
    backup = backup_dir / args.docx.name
    staged = backup_dir / f"{args.docx.stem}_staged.docx"
    shutil.copy2(args.docx, backup)
    document.save(staged)

    check = Document(staged)
    after = structure(check)
    if before != after:
        raise RuntimeError(f"Document structure changed: before={before}, after={after}")
    for item in items:
        index = int(item["paragraph_index"])
        if check.paragraphs[index].text.strip() != item["new_text"].strip():
            raise RuntimeError(f"Saved text mismatch at paragraph {index}")

    os.replace(staged, args.docx)
    result = {
        "document": str(args.docx.resolve()),
        "backup": str(backup),
        "before": before,
        "after": after,
        "rewrite_count": len(items),
        "items": audit_items,
    }
    args.audit.parent.mkdir(parents=True, exist_ok=True)
    args.audit.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({key: value for key, value in result.items() if key != "items"}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
