from __future__ import annotations

import argparse
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


CITATION_RE = re.compile(r"\[(\d+(?:(?:-|,|，)\d+)*)\]")
CAPTION_RE = re.compile(r"^(图|表)\s*(\d+)\.(\d+)\s*(.*)$")
REFERENCE_RE = re.compile(r"^\[(\d+)\]\s*(.*)$")


def pt_value(value):
    return None if value is None else round(value.pt, 2)


def font_values(rpr):
    fonts = rpr.rFonts if rpr is not None else None
    if fonts is None:
        return {"eastAsia": None, "ascii": None, "hAnsi": None, "cs": None}
    return {key: fonts.get(qn(f"w:{key}")) for key in ("eastAsia", "ascii", "hAnsi", "cs")}


def style_record(style):
    pf = style.paragraph_format
    return {
        "name": style.name,
        "base_style": style.base_style.name if style.base_style else None,
        "fonts": font_values(style.element.rPr),
        "size_pt": pt_value(style.font.size),
        "bold": style.font.bold,
        "alignment": int(pf.alignment) if pf.alignment is not None else None,
        "first_line_indent_pt": pt_value(pf.first_line_indent),
        "left_indent_pt": pt_value(pf.left_indent),
        "space_before_pt": pt_value(pf.space_before),
        "space_after_pt": pt_value(pf.space_after),
        "line_spacing": round(float(pf.line_spacing), 3) if isinstance(pf.line_spacing, float) else pt_value(pf.line_spacing),
        "line_spacing_rule": int(pf.line_spacing_rule) if pf.line_spacing_rule is not None else None,
        "keep_with_next": pf.keep_with_next,
        "keep_together": pf.keep_together,
    }


def all_paragraphs(document):
    for index, paragraph in enumerate(document.paragraphs):
        yield f"p:{index}", paragraph
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    yield f"t:{table_index}:{row_index}:{cell_index}:{paragraph_index}", paragraph


def expand_citation(token: str):
    result = set()
    for part in re.split(r"[,，]", token):
        if "-" in part:
            start, end = map(int, part.split("-", 1))
            result.update(range(start, end + 1))
        else:
            result.add(int(part))
    return result


def citation_is_superscript(paragraph, start, end):
    position = 0
    states = []
    for run in paragraph.runs:
        run_start, run_end = position, position + len(run.text)
        if max(start, run_start) < min(end, run_end):
            states.append(bool(run.font.superscript))
        position = run_end
    return bool(states) and all(states)


def border_summary(table):
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    result = {}
    if borders is not None:
        for edge in ("top", "insideH", "bottom", "left", "insideV", "right"):
            node = borders.find(qn(f"w:{edge}"))
            if node is not None:
                result[edge] = {
                    "val": node.get(qn("w:val")),
                    "size": node.get(qn("w:sz")),
                    "color": node.get(qn("w:color")),
                }
    return result


def main():
    parser = argparse.ArgumentParser(description="审计中文毕业论文 DOCX 的结构与常见格式问题。")
    parser.add_argument("docx", type=Path, help="待审计的论文 DOCX")
    parser.add_argument("--output", type=Path, required=True, help="UTF-8 JSON 审计结果")
    args = parser.parse_args()
    if not args.docx.is_file() or args.docx.suffix.lower() != ".docx":
        parser.error("输入必须是存在的 .docx 文件")

    document = Document(args.docx)
    used_styles = Counter(p.style.name for _, p in all_paragraphs(document))
    styles = {}
    for name in sorted(used_styles):
        try:
            styles[name] = style_record(document.styles[name])
        except KeyError:
            pass

    sections = []
    for index, section in enumerate(document.sections):
        sections.append({
            "index": index + 1,
            "width_cm": round(section.page_width.cm, 2),
            "height_cm": round(section.page_height.cm, 2),
            "top_cm": round(section.top_margin.cm, 2),
            "bottom_cm": round(section.bottom_margin.cm, 2),
            "left_cm": round(section.left_margin.cm, 2),
            "right_cm": round(section.right_margin.cm, 2),
            "header_cm": round(section.header_distance.cm, 2),
            "footer_cm": round(section.footer_distance.cm, 2),
            "different_first_page": section.different_first_page_header_footer,
            "header_linked": section.header.is_linked_to_previous,
            "footer_linked": section.footer.is_linked_to_previous,
            "header_text": " | ".join(p.text for p in section.header.paragraphs if p.text.strip()),
            "footer_text": " | ".join(p.text for p in section.footer.paragraphs if p.text.strip()),
        })

    direct_font_overrides = []
    direct_size_overrides = []
    for locator, paragraph in all_paragraphs(document):
        for run_index, run in enumerate(paragraph.runs):
            if not run.text.strip():
                continue
            fonts = font_values(run._r.rPr)
            if any(fonts.values()):
                direct_font_overrides.append({"locator": locator, "run": run_index, "style": paragraph.style.name, "text": run.text[:80], "fonts": fonts})
            if run.font.size is not None:
                direct_size_overrides.append({"locator": locator, "run": run_index, "style": paragraph.style.name, "text": run.text[:80], "size_pt": round(run.font.size.pt, 2)})

    reference_start = next(
        (i for i, p in enumerate(document.paragraphs) if p.style.name == "参考文献标题" or p.text.strip() == "参考文献"),
        len(document.paragraphs),
    )
    citations = []
    cited_numbers = set()
    for index, paragraph in enumerate(document.paragraphs[:reference_start]):
        for match in CITATION_RE.finditer(paragraph.text):
            numbers = expand_citation(match.group(1))
            cited_numbers.update(numbers)
            citations.append({
                "paragraph_index": index,
                "text": match.group(0),
                "numbers": sorted(numbers),
                "superscript": citation_is_superscript(paragraph, match.start(), match.end()),
                "context": paragraph.text[max(0, match.start() - 45):match.end() + 45],
            })

    references = []
    for index, paragraph in enumerate(document.paragraphs[reference_start + 1:], start=reference_start + 1):
        match = REFERENCE_RE.match(paragraph.text.strip())
        if match:
            references.append({"paragraph_index": index, "number": int(match.group(1)), "text": paragraph.text.strip(), "style": paragraph.style.name})
    reference_numbers = {item["number"] for item in references}

    captions = []
    for index, paragraph in enumerate(document.paragraphs):
        match = CAPTION_RE.match(paragraph.text.strip())
        if match:
            captions.append({
                "paragraph_index": index,
                "kind": match.group(1),
                "chapter": int(match.group(2)),
                "number": int(match.group(3)),
                "title": match.group(4),
                "style": paragraph.style.name,
                "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
            })

    caption_sequences = {}
    for kind in ("图", "表"):
        by_chapter = defaultdict(list)
        for item in captions:
            if item["kind"] == kind:
                by_chapter[item["chapter"]].append(item["number"])
        caption_sequences[kind] = {str(chapter): numbers for chapter, numbers in sorted(by_chapter.items())}

    table_records = []
    for table_index, table in enumerate(document.tables):
        font_sizes = Counter()
        alignments = Counter()
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    alignments[str(int(paragraph.alignment)) if paragraph.alignment is not None else "inherit"] += 1
                    for run in paragraph.runs:
                        if run.text.strip() and run.font.size is not None:
                            font_sizes[str(round(run.font.size.pt, 2))] += len(run.text)
        table_records.append({
            "index": table_index + 1,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "style": table.style.name if table.style else None,
            "font_sizes_by_chars": dict(font_sizes),
            "paragraph_alignments": dict(alignments),
            "borders": border_summary(table),
        })

    headings = [{"index": i, "style": p.style.name, "text": p.text.strip()} for i, p in enumerate(document.paragraphs) if p.style.name.startswith("Heading") or p.style.name.startswith("标题")]
    toc = [{"index": i, "style": p.style.name, "text": p.text.strip()} for i, p in enumerate(document.paragraphs) if p.style.name.lower().startswith("toc") or p.style.name.startswith("目录")]
    reference_sequence = [item["number"] for item in references]

    result = {
        "document": str(args.docx.resolve()),
        "counts": {"paragraphs": len(document.paragraphs), "tables": len(document.tables), "images": len(document.inline_shapes), "sections": len(document.sections), "headings": len(headings), "toc_entries": len(toc), "citations": len(citations), "references": len(references), "captions": len(captions)},
        "sections": sections,
        "styles_used": dict(used_styles),
        "styles": styles,
        "headings": headings,
        "toc": toc,
        "citations": citations,
        "citation_issues": {"not_superscript": [item for item in citations if not item["superscript"]], "cited_missing_reference": sorted(cited_numbers - reference_numbers), "uncited_references": sorted(reference_numbers - cited_numbers)},
        "references": references,
        "reference_sequence_ok": reference_sequence == list(range(1, len(references) + 1)),
        "captions": captions,
        "caption_sequences": caption_sequences,
        "tables": table_records,
        "direct_font_overrides": direct_font_overrides,
        "direct_size_overrides": direct_size_overrides,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"counts": result["counts"], "sections": sections, "citation_issues": result["citation_issues"], "reference_sequence_ok": result["reference_sequence_ok"], "caption_sequences": caption_sequences, "direct_font_overrides": len(direct_font_overrides), "direct_size_overrides": len(direct_size_overrides)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
